from __future__ import annotations

import io
import hashlib
import re
import zlib
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import olefile
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".hwp", ".hwpx", ".docx", ".txt", ".md", ".csv"}
MAX_FILE_COUNT = 10
MAX_FILE_SIZE = 20 * 1024 * 1024
MAX_TOTAL_CHARACTERS = 200_000
CHUNK_SIZE = 12_000
MIN_PDF_TEXT_CHARACTERS = 100
MIN_PDF_CHARACTERS_PER_PAGE = 20
ARTIFACT_DEFINITIONS = (
    ("REQUIREMENTS_DEFINITION", "요구사항 정의서", ("요구사항 정의서", "요구사항정의서")),
    ("FUNCTION_SPECIFICATION", "기능 명세서", ("기능 명세서", "기능명세서")),
    ("MEETING_MINUTES", "회의록", ("회의록", "회의 기록")),
    ("TEST_RESULTS", "테스트 결과서", ("테스트 결과서", "테스트결과서", "시험 결과서")),
    ("WEEKLY_REPORT", "주간 보고서", ("주간 보고서", "주간보고서", "주간 보고")),
    ("FINAL_REPORT", "최종 보고서", ("최종 보고서", "최종보고서", "완료 보고서")),
    ("UI_DESIGN", "UI 설계서", ("ui 설계서", "화면 설계서", "화면설계서")),
    ("PROPOSAL", "제안서", ("제안서",)),
    ("RFP", "RFP", ("rfp", "제안요청서", "제안 요청서")),
    ("WBS", "WBS", ("wbs", "작업분해구조", "작업 분해 구조")),
    ("ERD", "ERD", ("erd", "개체관계도", "개체 관계도")),
)


class DocumentExtractionError(ValueError):
    pass


@dataclass(frozen=True)
class UploadedDocument:
    file_name: str
    content_type: str | None
    content: bytes
    document_id: int | None = None


@dataclass(frozen=True)
class ParsedPage:
    page_number: int | None
    text: str


@dataclass(frozen=True)
class ParsedDocument:
    file_name: str
    file_type: str
    text: str
    content: bytes
    processing_mode: str
    document_id: int | None = None
    pages: tuple[ParsedPage, ...] = ()


class PlanningDocumentService:
    def parse_documents(self, uploads: list[UploadedDocument]) -> list[ParsedDocument]:
        if not uploads:
            raise DocumentExtractionError("최소 한 개의 문서가 필요합니다.")
        if len(uploads) > MAX_FILE_COUNT:
            raise DocumentExtractionError(f"문서는 최대 {MAX_FILE_COUNT}개까지 업로드할 수 있습니다.")

        parsed = [self._parse_document(upload) for upload in uploads]
        total_characters = sum(len(document.text) for document in parsed)
        if total_characters > MAX_TOTAL_CHARACTERS:
            raise DocumentExtractionError(
                f"추출된 문서 텍스트는 총 {MAX_TOTAL_CHARACTERS:,}자를 초과할 수 없습니다."
            )
        return parsed

    def build_chunks(self, documents: list[ParsedDocument]) -> list[dict[str, Any]]:
        chunks: list[dict[str, Any]] = []
        for document_index, document in enumerate(documents, start=1):
            if document.processing_mode != "TEXT":
                continue
            pages = document.pages or (ParsedPage(page_number=None, text=document.text),)
            for page in pages:
                text_parts = self._split_text_with_offsets(page.text)
                for index, (text, start_offset, end_offset) in enumerate(
                    text_parts,
                    start=1,
                ):
                    page_token = page.page_number or "na"
                    document_token = (
                        str(document.document_id)
                        if document.document_id is not None
                        else self._anonymous_document_token(
                            document.file_name,
                            document.content,
                            document_index,
                        )
                    )
                    chunks.append({
                        "chunk_id": f"{document_token}:{page_token}:{index}",
                        "document_id": document.document_id,
                        "source_document": document.file_name,
                        "page_number": page.page_number,
                        "chunk_index": index,
                        "text": text,
                        "start_offset": start_offset,
                        "end_offset": end_offset,
                    })
        return chunks

    def fallback_extract(self, chunk: dict[str, Any]) -> dict[str, Any]:
        text = chunk["text"]
        project_info = {
            "project_name": self._field_value(text, ("프로젝트명", "사업명", "과업명")),
            "project_goal": self._field_value(text, ("프로젝트 목표", "사업 목표", "추진 목적")),
            "client_organization": self._field_value(text, ("발주기관", "고객사", "발주처")),
            "period_start": None,
            "period_end": None,
            "key_features": self._list_value(text, ("주요 기능", "구축 범위")),
            "required_artifacts": self._artifact_list(text, ("주요 산출물", "산출물")),
            "acceptance_conditions": self._condition_list_value(
                text, ("검수 조건", "검수 기준")
            ),
            "budget_contract_conditions": self._condition_list_value(
                text, ("예산/계약 조건", "사업비", "계약 조건")
            ),
            "security_privacy_conditions": self._condition_list_value(
                text, ("보안/개인정보 조건", "보안 조건", "개인정보 조건")
            ),
        }
        period_text = self._field_value(text, ("수행 기간", "사업 기간", "계약 기간")) or text
        dates = re.findall(r"(20\d{2})[.\-/년]\s*(\d{1,2})[.\-/월]\s*(\d{1,2})", period_text)
        normalized_dates = [f"{year}-{int(month):02d}-{int(day):02d}" for year, month, day in dates]
        if normalized_dates:
            project_info["period_start"] = normalized_dates[0]
        if len(normalized_dates) > 1:
            project_info["period_end"] = normalized_dates[1]

        requirements = []
        for line in self._meaningful_lines(text):
            if not self._looks_like_requirement(line):
                continue
            cleaned = re.sub(r"^(REQ[-_ ]?\d+|요구사항\s*\d*)\s*[:.\-]?\s*", "", line, flags=re.I)
            cleaned = re.sub(r"^(필수|선택|권고|high|medium|low)\s*[:.\-]?\s*", "", cleaned, flags=re.I)
            local_offset = text.find(line)
            evidence = []
            if local_offset >= 0:
                chunk_start = int(chunk.get("start_offset") or 0)
                quote_text = line[:500]
                evidence.append({
                    "document_id": chunk.get("document_id"),
                    "source_document": chunk["source_document"],
                    "page_number": chunk.get("page_number"),
                    "chunk_id": chunk.get("chunk_id", ""),
                    "quote_text": quote_text,
                    "start_offset": chunk_start + local_offset,
                    "end_offset": chunk_start + local_offset + len(quote_text),
                    "bounding_boxes": [],
                })
            requirements.append({
                "function_name": self._function_name(cleaned),
                "requirement_text": cleaned[:1000],
                "category": self._category(cleaned),
                "priority": self._priority(line),
                "acceptance_criteria": None,
                "due_date": None,
                "deliverable_name": None,
                "security_condition": cleaned[:500] if self._is_security_text(cleaned) else None,
                "source_document": chunk["source_document"],
                "source_excerpt": line[:500],
                "evidences": evidence,
            })
        return {"project_info": project_info, "requirements": requirements[:50]}

    def consolidate(self, partials: list[dict[str, Any]]) -> dict[str, Any]:
        scalar_fields = (
            "project_name", "project_goal", "client_organization", "period_start", "period_end"
        )
        list_fields = (
            "key_features", "acceptance_conditions",
            "budget_contract_conditions", "security_privacy_conditions",
        )
        condition_fields = {
            "acceptance_conditions",
            "budget_contract_conditions",
            "security_privacy_conditions",
        }
        project_info: dict[str, Any] = {field: None for field in scalar_fields}
        project_info.update({field: [] for field in list_fields})
        project_info["required_artifacts"] = []

        for partial in partials:
            source = partial.get("project_info") or {}
            for field in scalar_fields:
                if not project_info[field] and source.get(field):
                    project_info[field] = source[field]
            for field in list_fields:
                for value in source.get(field) or []:
                    value = (
                        self._normalize_condition(value)
                        if field in condition_fields
                        else str(value).strip()
                    )
                    if value and value not in project_info[field]:
                        project_info[field].append(value)
            for artifact in source.get("required_artifacts") or []:
                if hasattr(artifact, "model_dump"):
                    artifact = artifact.model_dump(mode="json")
                if not isinstance(artifact, dict):
                    continue
                artifact_type = str(artifact.get("artifact_type") or "").strip()
                artifact_name = str(artifact.get("artifact_name") or "").strip()
                if not artifact_type or not artifact_name:
                    continue
                if any(
                    existing["artifact_type"] == artifact_type
                    for existing in project_info["required_artifacts"]
                ):
                    continue
                project_info["required_artifacts"].append({
                    "artifact_type": artifact_type,
                    "artifact_name": artifact_name,
                    "required_version": str(artifact.get("required_version") or "1.0").strip(),
                })

        requirements = []
        requirement_by_key: dict[str, dict[str, Any]] = {}
        for partial in partials:
            for requirement in partial.get("requirements") or []:
                text = str(requirement.get("requirement_text") or "").strip()
                if not text:
                    continue
                key = re.sub(r"\W+", "", text).lower()
                if not key:
                    continue
                evidences = self._deduplicate_evidences(
                    requirement.get("evidences") or []
                )
                if key in requirement_by_key:
                    existing = requirement_by_key[key]
                    existing["evidences"] = self._deduplicate_evidences(
                        [*existing["evidences"], *evidences]
                    )
                    continue
                source_document = requirement.get("source_document") or "unknown"
                source_excerpt = requirement.get("source_excerpt")
                if evidences:
                    source_document = evidences[0]["source_document"]
                    source_excerpt = evidences[0]["quote_text"]
                consolidated_requirement = {
                    "requirement_id": len(requirements) + 1,
                    "function_name": requirement.get("function_name") or "공통",
                    "requirement_text": text,
                    "category": requirement.get("category") or "UNSPECIFIED",
                    "priority": requirement.get("priority") or "UNSPECIFIED",
                    "acceptance_criteria": requirement.get("acceptance_criteria"),
                    "due_date": requirement.get("due_date"),
                    "deliverable_name": requirement.get("deliverable_name"),
                    "security_condition": requirement.get("security_condition"),
                    "source_document": source_document,
                    "source_excerpt": source_excerpt,
                    "evidences": evidences,
                }
                requirements.append(consolidated_requirement)
                requirement_by_key[key] = consolidated_requirement
        return {"project_info": project_info, "requirement_candidates": requirements[:200]}

    def _parse_document(self, upload: UploadedDocument) -> ParsedDocument:
        file_name = Path(upload.file_name).name
        suffix = Path(file_name).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            raise DocumentExtractionError(f"지원하지 않는 파일 형식입니다: {suffix or '확장자 없음'} ({supported})")
        if not upload.content:
            raise DocumentExtractionError(f"빈 파일입니다: {file_name}")
        if len(upload.content) > MAX_FILE_SIZE:
            raise DocumentExtractionError(f"파일 크기는 20MB를 초과할 수 없습니다: {file_name}")

        extractors = {
            ".pdf": self._extract_pdf,
            ".hwp": self._extract_hwp,
            ".hwpx": self._extract_hwpx,
            ".docx": self._extract_docx,
            ".txt": self._extract_plain_text,
            ".md": self._extract_plain_text,
            ".csv": self._extract_plain_text,
        }
        try:
            if suffix == ".pdf":
                pages = tuple(
                    ParsedPage(page_number=index, text=self._clean_text(page_text))
                    for index, page_text in enumerate(
                        self._extract_pdf_pages(upload.content),
                        start=1,
                    )
                )
                text = "\n".join(page.text for page in pages).strip()
                page_count = len(pages)
                processing_mode = (
                    "TEXT" if self._has_sufficient_pdf_text(text, page_count) else "PDF_VISION"
                )
            else:
                text = self._clean_text(extractors[suffix](upload.content))
                processing_mode = "TEXT"
                pages = (ParsedPage(page_number=None, text=text),)
        except DocumentExtractionError:
            raise
        except Exception as exc:
            raise DocumentExtractionError(f"문서를 읽을 수 없습니다: {file_name}") from exc
        if not text and processing_mode == "TEXT":
            raise DocumentExtractionError(
                f"텍스트를 추출할 수 없습니다: {file_name}"
            )
        return ParsedDocument(
            file_name=file_name,
            file_type=suffix[1:].upper(),
            text=text,
            content=upload.content,
            processing_mode=processing_mode,
            document_id=upload.document_id,
            pages=pages,
        )

    def _extract_pdf(self, content: bytes) -> tuple[str, int]:
        pages = self._extract_pdf_pages(content)
        return "\n".join(pages), len(pages)

    def _extract_pdf_pages(self, content: bytes) -> list[str]:
        reader = PdfReader(io.BytesIO(content))
        return [page.extract_text() or "" for page in reader.pages]

    def _has_sufficient_pdf_text(self, text: str, page_count: int) -> bool:
        meaningful_characters = len(re.sub(r"\s+", "", text))
        required_characters = max(
            MIN_PDF_TEXT_CHARACTERS,
            page_count * MIN_PDF_CHARACTERS_PER_PAGE,
        )
        return meaningful_characters >= required_characters

    def _extract_docx(self, content: bytes) -> str:
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs)

    def _extract_hwpx(self, content: bytes) -> str:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            section_names = sorted(
                name for name in archive.namelist()
                if name.lower().startswith("contents/section") and name.lower().endswith(".xml")
            )
            texts = []
            for name in section_names:
                root = ElementTree.fromstring(archive.read(name))
                texts.extend(value.strip() for value in root.itertext() if value.strip())
            return "\n".join(texts)

    def _extract_hwp(self, content: bytes) -> str:
        with olefile.OleFileIO(io.BytesIO(content)) as hwp:
            if not hwp.exists("FileHeader") or not hwp.exists("BodyText"):
                raise DocumentExtractionError("올바른 HWP 문서가 아닙니다.")
            header = hwp.openstream("FileHeader").read()
            compressed = len(header) > 36 and bool(header[36] & 1)
            section_paths = sorted(
                (path for path in hwp.listdir() if len(path) == 2 and path[0] == "BodyText"),
                key=lambda path: path[1],
            )
            paragraphs = []
            for path in section_paths:
                data = hwp.openstream(path).read()
                if compressed:
                    data = zlib.decompress(data, -15)
                paragraphs.extend(self._hwp_text_records(data))
            return "\n".join(paragraphs)

    def _hwp_text_records(self, data: bytes) -> list[str]:
        offset = 0
        texts = []
        while offset + 4 <= len(data):
            header = int.from_bytes(data[offset:offset + 4], "little")
            offset += 4
            tag_id = header & 0x3FF
            size = (header >> 20) & 0xFFF
            if size == 0xFFF:
                if offset + 4 > len(data):
                    break
                size = int.from_bytes(data[offset:offset + 4], "little")
                offset += 4
            record = data[offset:offset + size]
            offset += size
            if tag_id == 67:
                decoded = record.decode("utf-16le", errors="ignore")
                decoded = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", decoded)
                if decoded.strip():
                    texts.append(decoded.strip())
        return texts

    def _extract_plain_text(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "cp949", "utf-16"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentExtractionError("텍스트 파일 인코딩은 UTF-8, CP949 또는 UTF-16이어야 합니다.")

    def _split_text(self, text: str) -> list[str]:
        return [part for part, _, _ in self._split_text_with_offsets(text)]

    def _split_text_with_offsets(self, text: str) -> list[tuple[str, int, int]]:
        chunks: list[tuple[str, int, int]] = []
        cursor = 0
        text_length = len(text)

        while cursor < text_length:
            while cursor < text_length and text[cursor].isspace():
                cursor += 1
            if cursor >= text_length:
                break

            hard_end = min(cursor + CHUNK_SIZE, text_length)
            end = hard_end
            if hard_end < text_length:
                line_end = text.rfind("\n", cursor, hard_end + 1)
                if line_end > cursor:
                    end = line_end

            while end > cursor and text[end - 1].isspace():
                end -= 1
            if end <= cursor:
                end = hard_end

            chunk_text = text[cursor:end]
            if chunk_text:
                chunks.append((chunk_text, cursor, end))
            cursor = max(end, hard_end if end == cursor else end)

        if not chunks:
            return [(text[:CHUNK_SIZE], 0, min(len(text), CHUNK_SIZE))]
        return chunks

    def _anonymous_document_token(
        self,
        file_name: str,
        content: bytes,
        document_index: int,
    ) -> str:
        digest = hashlib.sha256(
            file_name.casefold().encode("utf-8")
            + b"\0"
            + hashlib.sha256(content).digest()
            + b"\0"
            + str(document_index).encode("ascii")
        ).hexdigest()
        return f"file-{digest[:12]}"

    def _deduplicate_evidences(
        self,
        evidences: list[Any],
    ) -> list[dict[str, Any]]:
        deduplicated: list[dict[str, Any]] = []
        seen: set[tuple[Any, Any, str]] = set()
        for evidence in evidences:
            if hasattr(evidence, "model_dump"):
                evidence = evidence.model_dump(mode="json")
            if not isinstance(evidence, dict):
                continue
            quote_text = str(evidence.get("quote_text") or "").strip()
            chunk_id = str(evidence.get("chunk_id") or "").strip()
            source_document = str(
                evidence.get("source_document") or ""
            ).strip()
            if not quote_text or not chunk_id or not source_document:
                continue
            key = (
                evidence.get("document_id"),
                evidence.get("page_number"),
                re.sub(r"\s+", " ", quote_text).casefold(),
            )
            if key in seen:
                continue
            seen.add(key)
            deduplicated.append({
                "document_id": evidence.get("document_id"),
                "source_document": source_document,
                "page_number": evidence.get("page_number"),
                "chunk_id": chunk_id,
                "quote_text": quote_text,
                "start_offset": evidence.get("start_offset"),
                "end_offset": evidence.get("end_offset"),
                "bounding_boxes": evidence.get("bounding_boxes") or [],
            })
        return deduplicated

    def _clean_text(self, text: str) -> str:
        text = text.replace("\u00a0", " ").replace("\x00", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _field_value(self, text: str, labels: tuple[str, ...]) -> str | None:
        for label in labels:
            match = re.search(rf"^{re.escape(label)}\s*[:：]\s*(.+)$", text, re.MULTILINE | re.I)
            if match:
                return match.group(1).strip()[:1000]
        return None

    def _list_value(self, text: str, labels: tuple[str, ...]) -> list[str]:
        value = self._field_value(text, labels)
        if not value:
            return []
        return [item.strip() for item in re.split(r"[,;/·]", value) if item.strip()]

    def _condition_list_value(self, text: str, labels: tuple[str, ...]) -> list[str]:
        return [
            normalized
            for item in self._list_value(text, labels)
            if (normalized := self._normalize_condition(item))
        ]

    def _normalize_condition(self, value: Any) -> str:
        normalized = re.sub(r"\s+", " ", str(value)).strip(" \t-•.;。")
        normalized = re.sub(
            r"(\S+)(?:을|를)\s+제출한\s+후",
            r"\1 제출 후",
            normalized,
        )
        normalized = re.sub(
            r"(?:하여야|해야|되어야|돼야|이어야)\s*(?:한다|함)$",
            "",
            normalized,
        ).strip()
        normalized = re.sub(r"(?:으로|로)\s*한다$", "", normalized).strip()
        normalized = re.sub(r"한다$", "", normalized).strip()
        normalized = re.sub(
            r"(\S+)(?:을|를)\s+"
            r"(적용|수행|제출|지급|암호화|검수|확인|점검|통과|준수|제공|처리|관리|보관|삭제|제한|분리|구축|실시)$",
            r"\1 \2",
            normalized,
        )
        normalized = re.sub(r"^(총\s*사업비|사업비)(?:은|는)\s+", r"\1 ", normalized)
        return normalized.strip(" \t-•.;。")

    def _artifact_list(self, text: str, labels: tuple[str, ...]) -> list[dict[str, str]]:
        artifacts = []
        seen_types = set()
        for item in self._list_value(text, labels):
            lowered = item.lower()
            for artifact_type, artifact_name, keywords in ARTIFACT_DEFINITIONS:
                if artifact_type in seen_types:
                    continue
                if any(keyword in lowered for keyword in keywords):
                    artifacts.append({
                        "artifact_type": artifact_type,
                        "artifact_name": artifact_name,
                        "required_version": "1.0",
                    })
                    seen_types.add(artifact_type)
                    break
        return artifacts

    def _meaningful_lines(self, text: str) -> list[str]:
        return [line.strip(" \t-•") for line in text.splitlines() if len(line.strip()) >= 8]

    def _looks_like_requirement(self, line: str) -> bool:
        lowered = line.lower()
        return bool(
            re.match(r"^(req[-_ ]?\d+|요구사항\s*\d*)", lowered)
            or any(phrase in lowered for phrase in ("해야 한다", "하여야 한다", "지원해야", "제공해야", "shall ", "must "))
        )

    def _function_name(self, text: str) -> str:
        for separator in (":", "-", "–"):
            if separator in text:
                candidate = text.split(separator, 1)[0].strip()
                if 1 < len(candidate) <= 80:
                    return candidate
        return "공통"

    def _priority(self, text: str) -> str:
        lowered = text.lower()
        if any(value in lowered for value in ("필수", "high", "높음")) or self._has_priority_grade(text, "상"):
            return "HIGH"
        if any(value in lowered for value in ("medium", "보통")) or self._has_priority_grade(text, "중"):
            return "MEDIUM"
        if any(value in lowered for value in ("선택", "low", "낮음")) or self._has_priority_grade(text, "하"):
            return "LOW"
        return "UNSPECIFIED"

    def _category(self, text: str) -> str:
        lowered = text.lower()
        category_keywords = (
            ("SECURITY", (
                "보안", "개인정보", "암호화", "접근권한", "접근 권한", "인증", "인가",
                "취약점", "침해", "security", "privacy",
            )),
            ("INTERFACE", (
                "외부 시스템", "외부시스템", "인터페이스", "연계", "api", "webhook",
                "integration",
            )),
            ("DATA", (
                "데이터", "데이터베이스", "database", "db ", "저장", "이관", "마이그레이션",
                "정합성", "데이터 품질",
            )),
            ("NON_FUNCTIONAL", (
                "성능", "응답시간", "응답 시간", "처리량", "동시접속", "동시 접속", "가용성",
                "확장성", "호환성", "사용성", "품질", "performance", "availability",
            )),
            ("OPERATION", (
                "운영", "유지보수", "유지 보수", "모니터링", "배포", "장애 대응", "장애대응",
                "헬프데스크", "관제", "백업", "복구",
            )),
            ("PROJECT_MANAGEMENT", (
                "일정", "산출물", "교육", "주간 보고", "월간 보고", "진행 보고", "보고서 제출",
                "검수", "회의", "사업관리", "사업 관리", "수행 인력", "프로젝트 관리",
            )),
        )
        for category, keywords in category_keywords:
            if any(keyword in lowered for keyword in keywords):
                return category
        if self._looks_like_requirement(text):
            return "FUNCTIONAL"
        return "UNSPECIFIED"

    def _has_priority_grade(self, text: str, grade: str) -> bool:
        return bool(re.search(rf"(?:우선순위|등급)\s*[:：]?\s*{grade}(?:\s|$)", text, re.I))

    def _is_security_text(self, text: str) -> bool:
        return any(keyword in text.lower() for keyword in ("보안", "개인정보", "암호화", "접근권한", "인증"))
