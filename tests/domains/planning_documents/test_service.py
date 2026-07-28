import io
import unittest
from unittest.mock import patch

from docx import Document

from app.domains.planning_documents.document_parser import (
    DocumentExtractionError,
    ParsedDocument,
    PlanningDocumentService,
    UploadedDocument,
)


class PlanningDocumentServiceTest(unittest.TestCase):
    def setUp(self):
        self.service = PlanningDocumentService()

    def test_parse_documents_requires_at_least_one_file(self):
        with self.assertRaisesRegex(DocumentExtractionError, "최소 한 개"):
            self.service.parse_documents([])

    def test_parse_rejects_empty_and_unsupported_files(self):
        for upload in (
            UploadedDocument("empty.txt", "text/plain", b""),
            UploadedDocument("image.png", "image/png", b"content"),
        ):
            with self.subTest(file=upload.file_name), self.assertRaises(
                DocumentExtractionError
            ):
                self.service.parse_documents([upload])

    def test_plain_text_supports_cp949_and_sanitizes_file_name(self):
        document = self.service.parse_documents(
            [
                UploadedDocument(
                    "../../rfp.txt",
                    "text/plain",
                    "프로젝트명: 테스트".encode("cp949"),
                )
            ]
        )[0]

        self.assertEqual(document.file_name, "rfp.txt")
        self.assertEqual(document.file_type, "TXT")
        self.assertIn("프로젝트명", document.text)

    def test_markdown_and_csv_use_plain_text_extraction(self):
        documents = self.service.parse_documents(
            [
                UploadedDocument(
                    "notes.md",
                    "text/markdown",
                    b"# Project\nPlanning notes",
                ),
                UploadedDocument(
                    "requirements.csv",
                    "text/csv",
                    b"id,requirement\n1,Login",
                ),
            ]
        )

        self.assertEqual(
            [document.file_type for document in documents],
            ["MD", "CSV"],
        )
        self.assertIn("Planning notes", documents[0].text)
        self.assertIn("1,Login", documents[1].text)

    def test_hwp_parses_body_text_records(self):
        text_bytes = "HWP project requirement".encode("utf-16le")
        record_header = (67 | (len(text_bytes) << 20)).to_bytes(
            4,
            "little",
        )

        class FakeOleDocument:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc_value, traceback):
                return False

            def exists(self, path):
                return path in {"FileHeader", "BodyText"}

            def listdir(self):
                return [["BodyText", "Section0"]]

            def openstream(self, path):
                if path == "FileHeader":
                    return io.BytesIO(bytes(37))
                return io.BytesIO(record_header + text_bytes)

        with patch(
            "app.domains.planning_documents.document_parser.olefile.OleFileIO",
            return_value=FakeOleDocument(),
        ):
            document = self.service.parse_documents(
                [
                    UploadedDocument(
                        "proposal.hwp",
                        "application/x-hwp",
                        b"synthetic-ole-container",
                    )
                ]
            )[0]

        self.assertEqual(document.file_type, "HWP")
        self.assertEqual(document.text, "HWP project requirement")

    def test_docx_parses_paragraphs_and_table_cells(self):
        source = Document()
        source.add_paragraph("Project overview")
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Field"
        table.cell(0, 1).text = "Value"
        output = io.BytesIO()
        source.save(output)

        parsed = self.service.parse_documents(
            [
                UploadedDocument(
                    "proposal.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    output.getvalue(),
                )
            ]
        )[0]

        self.assertIn("Project overview", parsed.text)
        self.assertIn("Field | Value", parsed.text)

    def test_total_extracted_character_limit_is_enforced(self):
        with patch(
            "app.domains.planning_documents.document_parser.MAX_TOTAL_CHARACTERS",
            5,
        ):
            with self.assertRaisesRegex(DocumentExtractionError, "총 5자"):
                self.service.parse_documents(
                    [
                        UploadedDocument(
                            "one.txt", "text/plain", b"abcd"
                        ),
                        UploadedDocument(
                            "two.txt", "text/plain", b"efgh"
                        ),
                    ]
                )

    def test_build_chunks_skips_vision_documents_and_preserves_order(self):
        documents = [
            ParsedDocument("one.txt", "TXT", "first", b"first", "TEXT"),
            ParsedDocument("scan.pdf", "PDF", "", b"%PDF", "PDF_VISION"),
            ParsedDocument("two.txt", "TXT", "second", b"second", "TEXT"),
        ]

        chunks = self.service.build_chunks(documents)

        self.assertEqual(
            [chunk["source_document"] for chunk in chunks],
            ["one.txt", "two.txt"],
        )
        self.assertEqual([chunk["chunk_index"] for chunk in chunks], [1, 1])

    def test_long_paragraph_is_split_at_configured_chunk_size(self):
        with patch(
            "app.domains.planning_documents.document_parser.CHUNK_SIZE", 5
        ):
            chunks = self.service._split_text("abcdefghijk")
        self.assertEqual(chunks, ["abcde", "fghij", "k"])

    def test_consolidate_deduplicates_artifacts_and_requirements(self):
        partials = [
            {
                "project_info": {
                    "project_name": "AIPM",
                    "required_artifacts": [
                        {
                            "artifact_type": "ERD",
                            "artifact_name": "ERD",
                            "required_version": "1.10",
                        }
                    ],
                    "acceptance_conditions": ["Tests must pass"],
                },
                "requirements": [
                    {
                        "function_name": "Login",
                        "requirement_text": "Users can log in.",
                        "source_document": "one.txt",
                    }
                ],
            },
            {
                "project_info": {
                    "project_name": "Ignored later name",
                    "required_artifacts": [
                        {
                            "artifact_type": "ERD",
                            "artifact_name": "Duplicate ERD",
                        },
                        "invalid",
                    ],
                },
                "requirements": [
                    {
                        "function_name": "Login duplicate",
                        "requirement_text": "Users can log in!",
                        "source_document": "two.txt",
                    },
                    {"requirement_text": "  "},
                ],
            },
        ]

        result = self.service.consolidate(partials)

        self.assertEqual(result["project_info"]["project_name"], "AIPM")
        self.assertEqual(len(result["project_info"]["required_artifacts"]), 1)
        self.assertEqual(
            result["project_info"]["required_artifacts"][0][
                "required_version"
            ],
            "1.10",
        )
        self.assertEqual(len(result["requirement_candidates"]), 1)
        self.assertEqual(
            result["requirement_candidates"][0]["requirement_id"], 1
        )


if __name__ == "__main__":
    unittest.main()
